from docx import Document
from docx.document import Document as _Document
import re


def extract_docx_text(docx_path: str, max_length: int = 12000) -> str:
    try:
        doc: _Document = Document(docx_path)
        full_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        content = "\n".join(full_text)
        content = re.sub(r'\n{3,}', '\n\n', content)

        return content[:max_length] if len(content) > max_length else content

    except Exception as e:
        return f"[Docx读取失败] {str(e)}"


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        print(extract_docx_text(sys.argv[1]))
